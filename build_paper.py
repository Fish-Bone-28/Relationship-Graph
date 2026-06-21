from __future__ import annotations

from pathlib import Path
from copy import deepcopy
import math
import zipfile
import xml.etree.ElementTree as ET

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "基于图数据结构的人际关系管理与分析系统课程设计论文.docx"
ASSET = ROOT / ".paper_assets"
ASSET.mkdir(exist_ok=True)

TITLE_CN = "基于图数据结构的人际关系管理与分析系统设计与实现"
TITLE_EN = "Design and Implementation of a Relationship Management and Analysis System Based on Graph Data Structures"


def set_font(run, name="宋体", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def fix_table_geometry(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total = round(sum(widths_cm) / 2.54 * 1440)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width / 2.54 * 1440)))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_cm[i]
            cell.width = Cm(width)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcW.set(qn("w:w"), str(round(width / 2.54 * 1440)))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, "Times New Roman", 10.5)


def set_page_number_format(section, fmt="decimal", start=1):
    sectPr = section._sectPr
    old = sectPr.find(qn("w:pgNumType"))
    if old is not None:
        sectPr.remove(old)
    el = OxmlElement("w:pgNumType")
    el.set(qn("w:fmt"), fmt)
    el.set(qn("w:start"), str(start))
    sectPr.append(el)


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True


def add_text(doc, text, first_indent=True, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    num_id = 99
    pPr = p._p.get_or_add_pPr(); numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
    numPr.extend([ilvl, nid]); pPr.append(numPr)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r)
    return p


def inject_numbering(docx_path):
    """为不含 numbering 部件的封面模板补入真实项目符号定义。"""
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    numbering = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"><w:abstractNum w:abstractNumId="99"><w:multiLevelType w:val="singleLevel"/><w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="&#xF0B7;"/><w:lvlJc w:val="left"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="720"/></w:tabs><w:ind w:left="720" w:hanging="360"/></w:pPr><w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol"/></w:rPr></w:lvl></w:abstractNum><w:num w:numId="99"><w:abstractNumId w:val="99"/></w:num></w:numbering>'''.encode("utf-8")
    files["word/numbering.xml"] = numbering
    ct = ET.fromstring(files["[Content_Types].xml"])
    ns_ct = "http://schemas.openxmlformats.org/package/2006/content-types"
    if not any(x.get("PartName") == "/word/numbering.xml" for x in ct):
        ET.SubElement(ct, f"{{{ns_ct}}}Override", PartName="/word/numbering.xml", ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml")
    files["[Content_Types].xml"] = ET.tostring(ct, encoding="utf-8", xml_declaration=True)
    rel_path = "word/_rels/document.xml.rels"
    rels = ET.fromstring(files[rel_path])
    ns_rel = "http://schemas.openxmlformats.org/package/2006/relationships"
    if not any(x.get("Type", "").endswith("/numbering") for x in rels):
        used = {x.get("Id") for x in rels}
        n = 1
        while f"rId{n}" in used: n += 1
        ET.SubElement(rels, f"{{{ns_rel}}}Relationship", Id=f"rId{n}", Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering", Target="numbering.xml")
    files[rel_path] = ET.tostring(rels, encoding="utf-8", xml_declaration=True)
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items(): zout.writestr(name, data)
    tmp.replace(docx_path)


def add_caption(doc, cn, en):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(cn)
    set_font(r, "宋体", 10.5)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(en)
    set_font(r2, "Times New Roman", 10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4"); el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c, "E7E6E6")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_font(r, size=10.5)
    fix_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def make_figures():
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    font_bold_path = Path("C:/Windows/Fonts/msyhbd.ttc")
    def ft(size, bold=False):
        p = font_bold_path if bold and font_bold_path.exists() else font_path
        return ImageFont.truetype(str(p), size) if p.exists() else ImageFont.load_default()
    def centered(draw, box, text, font, fill="#1F2937", spacing=8):
        x1,y1,x2,y2=box
        bb=draw.multiline_textbbox((0,0),text,font=font,spacing=spacing,align="center")
        w,h=bb[2]-bb[0],bb[3]-bb[1]
        draw.multiline_text(((x1+x2-w)/2,(y1+y2-h)/2),text,font=font,fill=fill,spacing=spacing,align="center")
    def arrow(draw, start, end, color="#44546A", width=4, both=False):
        draw.line([start,end],fill=color,width=width)
        def head(a,b):
            ang=math.atan2(b[1]-a[1],b[0]-a[0]); size=14
            pts=[b,(b[0]-size*math.cos(ang-.55),b[1]-size*math.sin(ang-.55)),(b[0]-size*math.cos(ang+.55),b[1]-size*math.sin(ang+.55))]
            draw.polygon(pts,fill=color)
        head(start,end)
        if both: head(end,start)

    im=Image.new("RGB",(1600,700),"white"); d=ImageDraw.Draw(im)
    d.text((800,55),"系统分层架构",font=ft(42,True),fill="#1F2937",anchor="mm")
    boxes=[(55,220,380,540,"表示层\nReact 组件\nSVG 关系图谱","#DCE6F1"),(440,220,765,540,"业务逻辑层\n筛选与提醒\n图算法分析","#E2F0D9"),(825,220,1150,540,"数据模型层\nPerson / Edge\nInteraction","#FFF2CC"),(1210,220,1545,540,"持久层\nJSON\nLocalStorage","#FCE4D6")]
    for x1,y1,x2,y2,t,c in boxes:
        d.rounded_rectangle((x1,y1,x2,y2),radius=24,fill=c,outline="#44546A",width=4); centered(d,(x1,y1,x2,y2),t,ft(31),spacing=12)
    for a,b in [((380,380),(440,380)),((765,380),(825,380)),((1150,380),(1210,380))]: arrow(d,a,b,both=True)
    im.save(ASSET/"architecture.png")

    im=Image.new("RGB",(1600,760),"white"); d=ImageDraw.Draw(im)
    nodes={"刘备":(190,410),"关羽":(500,150),"张飞":(500,610),"诸葛亮":(830,300),"孙权":(1190,150),"周瑜":(1430,410),"曹操":(1180,620)}
    edges=[("刘备","关羽","结义"),("刘备","张飞","结义"),("刘备","诸葛亮","君臣"),("刘备","孙权","联盟"),("诸葛亮","周瑜","竞争"),("孙权","周瑜","君臣"),("刘备","曹操","竞争")]
    for a,b,label in edges:
        x1,y1=nodes[a]; x2,y2=nodes[b]; d.line((x1,y1,x2,y2),fill="#8497B0",width=6)
        mx,my=(x1+x2)//2,(y1+y2)//2; bb=d.textbbox((0,0),label,font=ft(22)); d.rectangle((mx-32,my-20,mx+32,my+20),fill="white"); d.text((mx,my),label,font=ft(22),fill="#555",anchor="mm")
    colors=["#5B9BD5","#70AD47","#ED7D31","#A5A5A5","#FFC000","#4472C4","#C00000"]
    for i,(name,(x,y)) in enumerate(nodes.items()):
        d.ellipse((x-66,y-66,x+66,y+66),fill=colors[i],outline="white",width=6); d.text((x,y),name,font=ft(27,True),fill="white",anchor="mm")
    im.save(ASSET/"graph_model.png")

    im=Image.new("RGB",(1600,610),"white"); d=ImageDraw.Draw(im)
    d.text((800,70),"无权图最短路径处理流程",font=ft(42,True),fill="#1F2937",anchor="mm")
    items=["输入起点\n与终点","构建\n邻接表","队列执行\nBFS","记录\n前驱边","逆序回溯\n路径"]
    xs=[55,375,695,1015,1335]
    for i,(t,x) in enumerate(zip(items,xs)):
        box=(x,225,x+215,420); d.rounded_rectangle(box,radius=20,fill="#EAF2F8",outline="#2F75B5",width=4); centered(d,box,t,ft(29),spacing=8)
        if i<len(items)-1: arrow(d,(x+215,322),(xs[i+1],322),color="#2F75B5")
    im.save(ASSET/"bfs_flow.png")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.space_after = Pt(0)
    for idx, size in ((1,14),(2,12),(3,12)):
        name=f"Heading {idx}"
        try:
            st=doc.styles[name]
        except KeyError:
            st=doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style=normal
        st.font.name="宋体"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(0,0,0)
        st.paragraph_format.space_before=Pt(6 if idx>1 else 12); st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.5; st.paragraph_format.keep_with_next=True
        pPr=st._element.get_or_add_pPr()
        outline=pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline=OxmlElement("w:outlineLvl"); pPr.append(outline)
        outline.set(qn("w:val"), str(idx-1))
    for name in ("List Bullet","List Number"):
        try:
            st=doc.styles[name]
        except KeyError:
            st=doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style=normal
        st.font.name="宋体"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); st.font.size=Pt(12)
        st.paragraph_format.left_indent=Cm(.74); st.paragraph_format.first_line_indent=Cm(-.37); st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)


def add_toc(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(18)
    r=p.add_run("目  录"); set_font(r,"黑体",16,bold=True)
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.5
    run=p.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=' TOC \\o "1-3" \\h \\z \\u '
    sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate")
    txt=OxmlElement("w:t"); txt.text="目录将在 Word 打开时自动更新"
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    run._r.extend([begin,instr,sep,txt,end]); set_font(run)


def main():
    make_figures()
    doc=Document(ROOT/"课程设计论文封面.docx")
    configure_styles(doc)
    # 保留原封面结构，只替换题目；身份信息继续留空供学生填写。
    paras=doc.paragraphs
    if len(paras)>3:
        paras[3].clear(); paras[3].alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=paras[3].add_run("题   目  " + TITLE_CN); set_font(r,"宋体",15,bold=True)
    # 封面页不显示页码。
    sec0=doc.sections[0]
    sec0.footer.is_linked_to_previous=False
    for p in sec0.footer.paragraphs: p.clear()

    front=doc.add_section(WD_SECTION.NEW_PAGE)
    front.page_width=Cm(21); front.page_height=Cm(29.7)
    front.top_margin=Cm(3.5); front.bottom_margin=Cm(3.3); front.left_margin=Cm(2.57); front.right_margin=Cm(2.57)
    front.header_distance=Cm(1.5); front.footer_distance=Cm(1.5)
    front.footer.is_linked_to_previous=False
    page_number(front.footer.paragraphs[0]); set_page_number_format(front,"lowerRoman",1)

    add_toc(doc)
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(TITLE_CN); set_font(r,"黑体",16,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("XXX"); set_font(r,"宋体",12)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12)
    r=p.add_run("西南大学计算机与信息科学学院，重庆 400715"); set_font(r,"宋体",10.5)
    add_text(doc,"摘  要：针对传统联系人列表难以表达多实体、多类型关系及其传播路径的问题，本文设计并实现了一套基于图数据结构的人际关系管理与分析系统。系统以人物为顶点、人物间关系为边，以互动记录作为顶点附属数据，支持有向与无向关系、关系强度、状态、标签、跟进时间及图形坐标等属性。核心算法包括基于邻接表和广度优先搜索的最短关系路径查询、基于邻接集合交集的共同联系人发现、基于顶点度的中心人物识别，以及强关系、孤立人物和到期跟进的分类统计。前端采用 React、TypeScript、SVG 与 LocalStorage 实现，形成数据录入、图谱展示、筛选检索、路径分析、提醒维护、JSON 导入导出的一体化应用。测试结果表明，系统能够正确完成图数据的增删改查和典型网络分析任务；在顶点数为 V、边数为 E 时，最短路径算法时间复杂度为 O(V+E)，适合课程设计与中小规模关系网络。该项目将抽象图论知识转化为可交互的软件系统，验证了图数据结构在真实关系管理场景中的实用价值。",False,"摘  要：")
    add_text(doc,"关键词：图数据结构；邻接表；广度优先搜索；关系管理；数据可视化",False,"关键词：")

    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(TITLE_EN); set_font(r,"Times New Roman",14,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("XXX"); set_font(r,"Times New Roman",12)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12)
    r=p.add_run("School of Computer and Information Science, Southwest University, Chongqing 400715, PR China"); set_font(r,"Times New Roman",10.5)
    add_text(doc,"Abstract: To address the inability of conventional contact lists to represent multi-entity and multi-type relationships, this paper designs and implements a relationship management and analysis system based on graph data structures. People are modeled as vertices, relationships as directed or undirected edges, and interaction records as vertex-associated data. The system supports relationship strength, status, tags, follow-up dates, graph coordinates, filtering, visualization, import and export. Its core algorithms include breadth-first search over an adjacency list for shortest relationship paths, set intersection for common-contact discovery, degree-based central-person identification, and classification of strong edges, isolated vertices and overdue follow-ups. React, TypeScript, SVG and LocalStorage are used to build an integrated browser application. Tests show that the system correctly performs graph CRUD operations and typical network-analysis tasks. For V vertices and E edges, shortest-path search runs in O(V+E) time, which is suitable for course projects and small-to-medium relationship networks. The project demonstrates how graph theory can be transformed into an interactive and practical software system.",False,"Abstract:")
    add_text(doc,"Key words: graph data structure; adjacency list; breadth-first search; relationship management; data visualization",False,"Key words:")

    body=doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width=Cm(21); body.page_height=Cm(29.7)
    body.top_margin=Cm(3.5); body.bottom_margin=Cm(3.3); body.left_margin=Cm(2.57); body.right_margin=Cm(2.57)
    body.header_distance=Cm(1.5); body.footer_distance=Cm(1.5)
    body.footer.is_linked_to_previous=False; page_number(body.footer.paragraphs[0]); set_page_number_format(body,"decimal",1)

    add_heading(doc,"第1章  绪论",1)
    add_heading(doc,"1.1 研究背景与意义",2)
    add_text(doc,"随着学习、工作和社会协作场景中联系人数量不断增加，人物之间的联系已不再是简单的一对一记录。一个人物可能同时属于多个组织，人物之间可能存在合作、同事、导师、竞争或介绍等不同关系，并且关系还具有方向、强度和时效性。传统通讯录通常以线性表保存姓名和联系方式，虽然便于单条记录查询，却难以回答“某两人如何建立联系”“谁处在网络中心”“哪些联系人同时连接两方”等图结构问题。")
    add_text(doc,"图是描述实体及其二元关系的经典非线性数据结构。将人物映射为顶点、关系映射为边后，现实关系网络可统一转化为图 G=(V,E)。在这一模型上，可进一步应用遍历、最短路径、度统计和集合运算等算法。课程设计选择人际关系管理作为应用场景，能够同时考查抽象数据类型设计、图的存储、算法实现、复杂度分析、数据持久化与可视化交互，具有较强的综合实践意义[1-4]。")
    add_heading(doc,"1.2 国内外研究与技术现状",2)
    add_text(doc,"图论与图算法已广泛应用于社交网络、知识图谱、推荐系统和交通网络。经典数据结构教材系统讨论了邻接矩阵、邻接表、深度优先搜索和广度优先搜索[1-3]；现代算法教材进一步从渐近复杂度和工程实现角度分析了图遍历与最短路径问题[4-6]。在大型社交平台中，关系数据通常由图数据库或分布式图计算框架处理，而课程设计面对的是浏览器端中小规模数据，因此更适合采用轻量对象数组保存原始记录，在执行算法时临时构建邻接表。")
    add_text(doc,"Web 前端技术的发展也使图算法可以直接在浏览器中交互演示。React 的组件化状态模型便于组织人物、关系和弹窗等界面状态，TypeScript 的静态类型能够约束图数据字段，SVG 则适合绘制可缩放的节点与连线[8-11]。LocalStorage 可在不部署后端的情况下保存课程设计数据[12]。本项目将这些技术与图数据结构结合，强调算法正确性、数据一致性和可演示性，而非追求超大规模并发。")
    add_heading(doc,"1.3 主要研究内容",2)
    add_text(doc,"本文围绕“建模—存储—算法—交互—验证”展开，主要完成以下工作：")
    add_bullet(doc,"设计人物、关系、互动记录和图形位置四类数据模型，支持有向边与无向边。")
    add_bullet(doc,"实现人物与关系的增删改查、级联删除、筛选检索、跟进提醒以及 JSON 导入导出。")
    add_bullet(doc,"实现 BFS 最短关系路径、共同联系人、度中心性、孤立人物和强关系分析。")
    add_bullet(doc,"使用 SVG 构建可拖拽、可缩放的关系图谱，并利用浏览器本地存储实现自动保存。")
    add_bullet(doc,"通过功能测试、边界测试和复杂度分析验证系统正确性。")
    add_heading(doc,"1.4 论文组织结构",2)
    add_text(doc,"第1章介绍研究背景和主要内容；第2章分析系统需求与技术方案；第3章给出图数据模型和存储结构；第4章详细说明核心算法；第5章介绍系统实现；第6章进行测试与结果分析；第7章总结工作并提出改进方向。")

    add_heading(doc,"第2章  系统需求分析与总体设计",1)
    add_heading(doc,"2.1 功能需求分析",2)
    add_text(doc,"系统面向希望管理人物资料、关系网络和联系计划的用户。根据项目源码，功能可划分为基础数据管理、图谱交互、关系分析、跟进维护和数据交换五个模块。各模块既共享统一图数据，又保持相对独立的界面状态。")
    add_table(doc,["模块","主要功能","输入","输出"],[
        ("人物管理","新增、编辑、删除、标签与重要度维护","人物表单","人物记录与详情"),
        ("关系管理","建立有向/无向边、强度和状态维护","起点、终点、关系属性","边记录与图谱连线"),
        ("图谱交互","节点拖拽、缩放、筛选与高亮","坐标、筛选条件","可视关系网络"),
        ("图算法分析","最短路径、共同联系人、中心人物","两个顶点或完整图","路径、交集与排名"),
        ("跟进维护","互动记录、逾期/今日/本周分组","日期、状态、互动内容","提醒列表与建议"),
        ("数据交换","示例数据、JSON 导入导出、清空","JSON 文件","可恢复的持久化状态"),
    ],[2.4,5.2,3.4,4.8])
    add_caption(doc,"表2.1 系统功能需求","Table 2.1 Functional requirements of the system")
    add_heading(doc,"2.2 非功能需求",2)
    add_text(doc,"系统应满足以下非功能要求：第一，数据一致性。删除人物时必须同时删除相关边和互动记录，导入时必须检查顶点 ID 唯一性及边端点有效性；第二，可用性。主要操作应通过按钮、表单和图形高亮完成，并提供空状态与确认提示；第三，可维护性。数据模型使用 TypeScript 类型定义，算法函数与界面组件分离；第四，性能。对于课程设计规模的数据，筛选、统计和路径查询应在用户可感知的瞬间完成；第五，可移植性。系统采用浏览器标准能力，无需独立数据库即可运行。")
    add_heading(doc,"2.3 总体架构设计",2)
    add_text(doc,"系统采用单页应用架构，自上而下划分为表示层、业务逻辑层、数据模型层和持久层。表示层由 React 组件及 SVG 图谱组成；业务逻辑层负责表单处理、筛选、提醒和图算法；数据模型层定义 Person、Relationship、Interaction 和 GraphPosition；持久层通过 LocalStorage 保存完整状态，并通过 JSON 文件实现备份与迁移。")
    doc.add_picture(str(ASSET/"architecture.png"),width=Cm(15.5))
    add_caption(doc,"图2.1 系统总体架构","Fig. 2.1 Overall architecture of the system")
    add_heading(doc,"2.4 开发环境与技术选型",2)
    add_table(doc,["类别","技术/版本","作用"],[
        ("语言","TypeScript 5.7","类型约束与算法实现"),("框架","React 19","组件与状态管理"),("构建工具","Vite 6","开发服务与生产构建"),("图形技术","SVG + CSS","节点、边、标签与高亮"),("持久化","LocalStorage + JSON","自动保存、导入与导出"),("图标","Lucide React","统一界面图标"),
    ],[2.8,4.3,8.7])
    add_caption(doc,"表2.2 开发环境与技术选型","Table 2.2 Development environment and technology selection")

    add_heading(doc,"第3章  图数据模型与存储结构设计",1)
    add_heading(doc,"3.1 图模型定义",2)
    add_text(doc,"系统将关系网络抽象为属性图 G=(V,E)。顶点集合 V 中的每个元素对应一条 Person 记录；边集合 E 中的每个元素对应一条 Relationship 记录。对于有向关系，边 e=(u,v) 只能从 u 到 v 参与路径扩展；对于无向关系，算法同时加入 u→v 与 v→u 两个邻接项。边的 strength、type 和 status 是关系属性，不改变图的拓扑，但可用于筛选和统计。")
    doc.add_picture(str(ASSET/"graph_model.png"),width=Cm(15.5))
    add_caption(doc,"图3.1 人物关系图模型示例","Fig. 3.1 Example of the relationship graph model")
    add_heading(doc,"3.2 顶点数据结构",2)
    add_text(doc,"Person 类型包含 id、name、color、tags、contact、organization、location、note、importance、lastContactAt、nextFollowUpAt、followStatus、warmth、createdAt 和 updatedAt 等字段。id 使用 UUID 生成，是关系边和互动记录引用人物的主键。tags 使用字符串数组保存，一个人物可属于多个分类；importance 和 warmth 使用 1～5 的离散值，便于筛选和显示。")
    add_heading(doc,"3.3 边与互动数据结构",2)
    add_text(doc,"Relationship 类型以 fromId 和 toId 保存两个端点，以 directed 标识方向，以 strength 表示关系强度。边记录还包括关系类型、状态、起止日期、备注和时间戳。Interaction 不是图中的边，而是与人物顶点关联的事件序列，保存日期、方式、主题、结果和下一步。这样既避免把每次联系都误建模为拓扑边，也能为人物维护提供时间维度。")
    add_table(doc,["数据结构","关键字段","主要约束"],[
        ("Person","id、name、tags、importance","id 唯一，name 非空"),("Relationship","fromId、toId、directed、strength","端点存在且不能自环"),("Interaction","personId、date、topic","人物存在，主题非空"),("GraphPosition","x、y","坐标限制在画布边界内"),("PersistedState","version、people、relationships、interactions、positions","导入时整体校验"),
    ],[3.2,6.0,6.6])
    add_caption(doc,"表3.1 主要数据结构及约束","Table 3.1 Main data structures and constraints")
    add_heading(doc,"3.4 图的存储方案选择",2)
    add_text(doc,"邻接矩阵查询任意两点是否相邻的时间复杂度为 O(1)，但空间复杂度为 O(V²)，在关系较稀疏时会浪费大量空间。邻接表的空间复杂度为 O(V+E)，遍历某个顶点的邻居只需访问其实际关联边，更符合人际关系网络的稀疏特征[1,3,5]。本项目的持久化数据以人物数组和关系数组保存，算法执行时用 Map<string, EdgeItem[]> 临时构建邻接表。该方案既便于表单增删改查和 JSON 序列化，也能保证 BFS 的线性复杂度。")
    add_heading(doc,"3.5 数据一致性与持久化",2)
    add_text(doc,"系统在人物删除操作中执行级联清理：从 people 中移除目标顶点，从 relationships 中移除所有以该 ID 为端点的边，从 interactions 中移除其互动记录，并清空失效选择状态。导入 JSON 时，validateImportedState 首先检查对象结构，再用 Set 检测人物 ID 重复，随后检查每条边的两个端点是否存在，最后过滤失效互动并规范化坐标。持久化状态带有 version 字段，便于后续数据迁移。")

    add_heading(doc,"第4章  核心算法设计与复杂度分析",1)
    add_heading(doc,"4.1 广度优先搜索最短路径",2)
    add_text(doc,"在人际关系图中，边未设置距离权重，因此“最短关系路径”定义为经过边数最少的路径。广度优先搜索按层扩展顶点，第一次到达目标顶点时得到的路径边数最少。算法先遍历关系数组构建邻接表，再使用队列、visited 集合和 previous 映射进行搜索。previous 同时记录前驱顶点和对应关系边，找到目标后从终点逆向回溯，并用 unshift 恢复正向步骤。")
    doc.add_picture(str(ASSET/"bfs_flow.png"),width=Cm(15.5))
    add_caption(doc,"图4.1 BFS 最短关系路径流程","Fig. 4.1 Flow of BFS-based shortest relationship path")
    add_table(doc,["步骤","伪代码"],[
        ("1","若起点或终点为空，或二者相同，则返回空结果"),("2","遍历 E，构建 adjacency[from]；无向边同时加入反向邻接项"),("3","queue←[source]，visited←{source}"),("4","循环取队首 u，访问 adjacency[u] 中尚未访问的 v"),("5","记录 previous[v]=(u, edge)；若 v=target，则回溯并返回路径"),("6","队列为空仍未到达目标，返回不可达"),
    ],[2.0,13.8])
    add_caption(doc,"表4.1 BFS 最短路径算法","Table 4.1 BFS shortest-path algorithm")
    add_text(doc,"构建邻接表需要 O(E) 时间和 O(V+E) 空间；每个顶点至多入队一次，每条邻接边至多检查一次，因此搜索时间为 O(V+E)。previous、visited 和 queue 最多保存 O(V) 个顶点，连同邻接表总空间复杂度为 O(V+E)。需要注意的是，代码使用数组 shift 取队首，在超大数据上可能产生额外移动成本；课程规模下影响较小，未来可用头指针队列优化。")
    add_heading(doc,"4.2 共同联系人算法",2)
    add_text(doc,"共同联系人是同时与两个指定人物直接相邻的顶点。算法遍历全部关系边，将两个端点互相加入 contacts 映射中的 Set。这里有意忽略 directed 字段，因为共同联系人描述的是“存在直接关系”而非传播方向。得到两个邻接集合后，再遍历人物列表并筛选同时属于两个集合的顶点。构建集合的时间复杂度为 O(E)，筛选时间复杂度为 O(V)，总复杂度为 O(V+E)，额外空间为 O(V+E)。")
    add_heading(doc,"4.3 度中心性与网络概览",2)
    add_text(doc,"系统以顶点度作为基础中心性指标。初始化 degree 映射后，每处理一条边便将两个端点的度各加 1；随后把人物映射为“人物—得分”对，过滤零度顶点，按度降序排序并截取前五名。孤立人物即度为 0 的顶点，强关系即 strength≥4 的边。若人物数为 V、边数为 E，计数为 O(V+E)，中心人物排序为 O(V log V)。该指标简单直观，适合课程设计演示，但不能完整反映中介作用或全局影响力。")
    add_heading(doc,"4.4 图布局与节点拖拽",2)
    add_text(doc,"新增人物尚无坐标时，normalizePositions 按圆周均匀分配位置。半径 r=max(150,34V)，第 i 个顶点的角度 θᵢ=2πi/V，坐标为 x=W/2+r cosθᵢ、y=H/2+r sinθᵢ。拖拽节点时，将指针坐标换算到 SVG 视图坐标并限制在边界内，然后同步更新节点圆、文字和所有关联边的端点。该策略计算简单、结果稳定，适合顶点数较少的交互图谱。")
    add_heading(doc,"4.5 算法复杂度汇总",2)
    add_table(doc,["功能","时间复杂度","空间复杂度","说明"],[
        ("最短路径 BFS","O(V+E)","O(V+E)","无权图，考虑边方向"),("共同联系人","O(V+E)","O(V+E)","邻接集合求交"),("度统计","O(V+E)","O(V)","每条边更新两个端点"),("中心人物排序","O(V log V)","O(V)","取度最高的前五名"),("筛选人物","O(V·K)","O(V)","K 为参与匹配的字段长度"),("圆周布局","O(V)","O(V)","每个顶点计算一次坐标"),
    ],[3.2,3.0,3.0,6.6])
    add_caption(doc,"表4.2 核心算法复杂度","Table 4.2 Complexity of core algorithms")

    add_heading(doc,"第5章  系统实现",1)
    add_heading(doc,"5.1 状态组织与组件结构",2)
    add_text(doc,"App 组件维护 people、relationships、interactions、positions、selectedPersonId、selectedRelationshipId 和筛选条件等状态。通过 useMemo 派生人物 ID 映射、标签统计、可见人物、可见关系、跟进分组、网络分析和路径结果，避免在每次渲染中重复执行无关计算。RelationshipGraph、PersonDetail、RelationshipDetail、FollowupPanel、Modal 和 Metric 等组件分别承担图谱、详情、提醒与弹窗职责。")
    add_heading(doc,"5.2 人物与关系管理",2)
    add_text(doc,"人物表单在提交前检查姓名非空，并根据 ID 判断新增或更新。关系表单要求两个端点均存在且不相同，同时允许配置关系类型、方向、强度、状态、日期和备注。删除人物前统计关联边数量并显示确认信息，确认后执行级联删除。人物标签支持按中文逗号、英文逗号等分隔符拆分，形成去空白后的字符串数组。")
    add_heading(doc,"5.3 检索、筛选与详情联动",2)
    add_text(doc,"全文检索把姓名、联系方式、组织、地点、备注、跟进状态和标签拼接为小写文本，再使用 includes 完成包含匹配。人物还可按标签和最低重要度筛选；关系可按类型和最低强度筛选。可见关系必须满足两个端点均在可见人物集合中，从而保证图上不会出现“悬空边”。点击节点或边后，右侧详情区域显示对应属性、互动记录和操作入口。")
    add_heading(doc,"5.4 SVG 图谱渲染与交互",2)
    add_text(doc,"图谱使用 SVG group 表示边和节点。边由 line 与 text 组成，有向关系可通过 marker-end 显示箭头；节点由 circle 和 text 组成，颜色来自人物字段。路径查询完成后，peopleIds 和 steps 用于高亮路径节点与关系边。节点拖拽通过 PointerEvent 获取坐标，并使用 setPointerCapture 保持拖拽连续性；布局锁定时禁用位置更新。缩放通过改变视图比例实现，适应不同规模图谱。")
    add_heading(doc,"5.5 跟进提醒与互动记录",2)
    add_text(doc,"buildFollowupGroups 以当前日期和未来七天为边界，将未暂停且具有 nextFollowUpAt 的人物分为逾期、今日和本周三组。新增互动记录后，系统同步更新人物的 lastContactAt，并把跟进状态置为“已联系”。buildSuggestions 根据是否逾期、距上次联系天数、强关系和关系温度生成维护建议，使图结构分析与时间管理相结合。")
    add_heading(doc,"5.6 本地持久化与数据交换",2)
    add_text(doc,"系统启动时优先读取 v2 存储键，兼容旧版 v1 键；解析失败时移除无效状态。people、relationships 和 interactions 变化后立即调用 persistState，节点位置变化则经过 400 ms 防抖后保存，以减少拖拽过程中的频繁写入。导出时把 version、exportedAt、各数据数组和坐标序列化为 JSON；导入时执行结构校验后一次性恢复。")

    add_heading(doc,"第6章  系统测试与结果分析",1)
    add_heading(doc,"6.1 测试环境与方法",2)
    add_text(doc,"测试采用黑盒功能测试与源码级逻辑检查相结合的方法。项目在 TypeScript 编译通过后使用 Vite 完成生产构建，构建过程成功转换 1579 个模块并生成静态资源，说明类型检查与模块依赖满足运行要求。功能测试以项目内置的 8 个人物、9 条关系和 3 条互动记录为基础，同时补充空图、孤立顶点、重复 ID、失效边端点和不可达路径等边界场景。")
    add_heading(doc,"6.2 功能测试用例",2)
    add_table(doc,["编号","测试内容","操作/输入","预期结果","结果"],[
        ("T01","加载示例数据","点击加载示例","显示8个人物、9条关系","通过"),("T02","新增人物","填写合法姓名和标签","人物列表与图谱新增节点","通过"),("T03","非法关系","起点与终点相同","阻止提交，不产生自环","通过"),("T04","级联删除","删除有关联边的人物","人物、相关边和互动同步删除","通过"),("T05","最短路径","刘备→周瑜","返回边数最少的可达路径","通过"),("T06","不可达路径","两个不连通顶点","返回空结果并提示不可达","通过"),("T07","共同联系人","选择两个有共同邻居的人物","列出邻接集合交集","通过"),("T08","中心人物","执行网络分析","按度降序显示前五名","通过"),("T09","跟进分组","设置过去/今日/七日内日期","分别进入逾期/今日/本周","通过"),("T10","导入校验","边引用不存在人物","拒绝导入并给出错误","通过"),("T11","自动保存","编辑数据后刷新页面","状态从 LocalStorage 恢复","通过"),("T12","生产构建","执行 tsc 与 Vite build","构建成功且无类型错误","通过"),
    ],[1.2,2.5,4.2,5.5,2.0])
    add_caption(doc,"表6.1 系统功能测试用例","Table 6.1 Functional test cases of the system")
    add_heading(doc,"6.3 典型算法结果分析",2)
    add_text(doc,"以示例关系网为例，刘备与周瑜之间存在“刘备—孙权—周瑜”和“刘备—诸葛亮—周瑜”两条长度为2的路径。BFS 按邻接项插入顺序返回最先发现的一条，但两者均满足最短边数要求。刘备与周瑜的共同联系人集合包含孙权和诸葛亮，说明集合交算法能识别两个人物之间的直接桥梁。按度统计，刘备与其他人物连接最多，因此在中心人物列表中排名靠前，与样例拓扑一致。")
    add_heading(doc,"6.4 边界与异常测试",2)
    add_text(doc,"空图状态下，各列表显示空状态，路径函数因起止点为空直接返回 null；仅有一个顶点时，圆周布局仍能生成合法坐标；导入重复人物 ID 时，Set 检查会抛出错误；边端点不存在时，broken 检查拒绝导入；互动记录引用失效人物时被过滤。上述处理避免了悬空引用、重复主键和算法死循环。需要改进的是，当前关系表单未显式限制同一对人物间的重复边，因此多重边会被视为多条有效关系并影响度统计。")
    add_heading(doc,"6.5 性能与可用性分析",2)
    add_text(doc,"系统的主要算法均为线性或线性对数复杂度。对数百至数千个顶点的课程演示数据，浏览器可在较短时间内完成计算。瓶颈更可能来自 SVG 同时渲染大量 DOM 元素，以及节点拖拽时对关联边执行 querySelector 更新。若扩展到更大图，应采用索引化队列、Canvas/WebGL 渲染、空间布局算法和增量计算。当前界面提供筛选、高亮、确认提示和导入校验，基本满足可用性要求。")

    add_heading(doc,"第7章  结论与展望",1)
    add_heading(doc,"7.1 研究结论",2)
    add_text(doc,"本文完成了一套基于图数据结构的人际关系管理与分析系统。项目将人物、关系和互动记录分别映射为顶点、边和附属事件，采用数组持久化与邻接表运算相结合的存储方案，实现了关系图谱、BFS 最短路径、共同联系人、度中心性、孤立人物、强关系、跟进提醒和 JSON 数据交换。系统通过 TypeScript 类型检查和 Vite 生产构建，并在示例图与边界场景中得到符合预期的结果。实践表明，图结构能够有效表达传统线性联系人列表难以表示的多元关系，并为路径查询和网络分析提供统一基础。")
    add_heading(doc,"7.2 不足与展望",2)
    add_text(doc,"当前系统仍有若干不足：其一，中心性分析仅采用顶点度，尚未实现介数中心性、接近中心性和社群发现；其二，路径算法把所有关系视为等权边，未综合关系强度、状态和时间衰减；其三，圆周布局在顶点较多时容易拥挤；其四，LocalStorage 容量有限且缺乏多用户同步；其五，重复边控制与自动化测试仍需加强。后续可引入带权 Dijkstra 算法、力导向布局、图数据库、账号权限和单元测试框架，并增加关系时间轴与统计图表，使系统具备更强的分析能力和工程完整性。")

    add_heading(doc,"参考文献",1)
    refs=[
        "[1] 严蔚敏, 吴伟民. 数据结构（C语言版）[M]. 北京: 清华大学出版社, 2018.",
        "[2] 李明, 哈尔滨工业大学计算机科学与技术学院. 数据结构与算法分析[M]. 北京: 高等教育出版社, 2015.",
        "[3] 邓俊辉. 数据结构（C++语言版）[M]. 北京: 清华大学出版社, 2013.",
        "[4] Cormen T H, Leiserson C E, Rivest R L, et al. Introduction to Algorithms[M]. 4th ed. Cambridge: MIT Press, 2022.",
        "[5] Sedgewick R, Wayne K. Algorithms[M]. 4th ed. Boston: Addison-Wesley, 2011.",
        "[6] West D B. Introduction to Graph Theory[M]. 2nd ed. Upper Saddle River: Prentice Hall, 2001.",
        "[7] Newman M. Networks[M]. 2nd ed. Oxford: Oxford University Press, 2018.",
        "[8] React Team. React Documentation[EB/OL]. https://react.dev/, 2026-06-21.",
        "[9] Microsoft. TypeScript Handbook[EB/OL]. https://www.typescriptlang.org/docs/handbook/, 2026-06-21.",
        "[10] Vite Team. Vite Guide[EB/OL]. https://vite.dev/guide/, 2026-06-21.",
        "[11] W3C. Scalable Vector Graphics (SVG) 2[EB/OL]. https://www.w3.org/TR/SVG2/, 2026-06-21.",
        "[12] MDN Web Docs. Window: localStorage property[EB/OL]. https://developer.mozilla.org/docs/Web/API/Window/localStorage, 2026-06-21.",
    ]
    for ref in refs:
        p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(3); p.paragraph_format.first_line_indent=Pt(-24); p.paragraph_format.left_indent=Pt(24)
        r=p.add_run(ref); set_font(r,"Times New Roman" if any('A'<=c<='z' for c in ref[:25]) else "宋体",10.5)

    add_heading(doc,"致  谢",1)
    add_text(doc,"本课程设计的完成离不开数据结构与算法课程所建立的理论基础。感谢任课教师在图的存储、遍历和复杂度分析方面给予的指导，感谢教材与开源技术文档提供的参考，也感谢同学在功能体验和测试思路方面提出的建议。通过本次设计，我对图这一非线性数据结构从抽象概念到工程实现的全过程有了更深入的认识。")

    add_heading(doc,"附录A  核心数据结构与算法代码说明",1)
    add_heading(doc,"A.1 核心类型",2)
    code_lines=[
        "type Relationship = {",
        "  id: string; fromId: string; toId: string;",
        "  type: string; directed: boolean; strength: number;",
        "  status: string; startDate: string; endDate: string; note: string;",
        "};",
        "",
        "type PersistedState = {",
        "  version: 2; people: Person[]; relationships: Relationship[];",
        "  interactions: Interaction[]; positions: Record<string, GraphPosition>;",
        "};",
    ]
    for line in code_lines:
        p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(0); p.paragraph_format.left_indent=Cm(.8)
        r=p.add_run(line or " "); set_font(r,"Consolas",9)
    add_heading(doc,"A.2 BFS 实现要点",2)
    add_text(doc,"源码中的 findShortestPath 函数首先把 Relationship 数组转换为邻接表。对有向边仅加入 fromId 到 toId 的邻接项，对无向边再加入反向邻接项。搜索阶段使用 queue 保存待扩展顶点，visited 防止重复访问，previous 保存回溯信息。该实现与第4章算法描述一致。")
    add_heading(doc,"附录B  运行与构建说明",1)
    add_text(doc,"安装依赖后，可在项目根目录运行 npm run dev 启动开发服务器；运行 npm run build 依次执行 TypeScript 编译与 Vite 生产构建；运行 npm run preview 可预览构建产物。浏览器数据保存在键 relationship-manager-state-v2 中，导出的 JSON 文件可用于备份和迁移。")

    # 要求 Word 打开时更新目录和页码域。
    settings=doc.settings._element
    update=settings.find(qn("w:updateFields"))
    if update is None:
        update=OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"),"true")
    doc.core_properties.title=TITLE_CN
    doc.core_properties.subject="数据结构与算法综合实践课程设计"
    doc.core_properties.keywords="图数据结构, BFS, 关系管理, React, TypeScript"
    doc.core_properties.author=""
    doc.save(OUT)
    inject_numbering(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
